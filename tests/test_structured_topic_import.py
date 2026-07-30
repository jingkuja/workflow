from __future__ import annotations

from io import BytesIO

from docx import Document
from sqlalchemy import select
from test_t2_service_flow import make_service, upload_file

from workflow.db.models import ContentProject, ImportBatch
from workflow.t2.contracts import StructuredTopicInput


def _arbitrary_document() -> bytes:
    document = Document()
    document.add_paragraph("本周 AI 市场观察")
    document.add_paragraph("人工智能代理正在进入企业流程，这是第一条候选选题的对应原文。")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "补充观察"
    table.cell(0, 1).text = "算力成本下降正在改变中小企业采用 AI 的节奏。"
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_structured_import_accepts_arbitrary_word_layout_and_deduplicates(tmp_path) -> None:
    service = make_service(tmp_path)
    content = _arbitrary_document()
    file_key = upload_file(
        service, actor_name="老板测试", role="BOSS", content=content
    )
    topics = [
        StructuredTopicInput(
            source_index="段落 2",
            title="AI 代理进入企业工作流",
            source_text="人工智能代理正在进入企业流程，这是第一条候选选题的对应原文。",
            script=None,
            confidence=0.92,
            evidence=["人工智能代理正在进入企业流程"],
        ),
        StructuredTopicInput(
            source_index="表格 1",
            title="算力降价推动中小企业采用 AI",
            source_text="算力成本下降正在改变中小企业采用 AI 的节奏。",
            confidence=0.65,
            evidence=["算力成本下降"],
        ),
    ]

    imported = service.import_structured_topics(
        actor_name="老板测试",
        original_filename="ChatGPT随手生成的选题.docx",
        idempotency_key="structured-import-0001",
        topics=topics,
        warnings=["第二条需要确认数据时效。"],
        schema_version="1.0",
        file_key=file_key,
        file_url=None,
    )

    data = imported["data"]
    assert data["created_count"] == 2
    assert data["parse_status"] == "COMPLETED"
    assert data["import_mode"] == "WORKBUDDY_STRUCTURED"
    assert data["schema_version"] == "1.0"
    assert data["warnings"] == ["第二条需要确认数据时效。"]

    with service.sessions() as session:
        projects = session.scalars(
            select(ContentProject).order_by(ContentProject.source_sequence)
        ).all()
        assert projects[0].source_content["script"] is None
        assert projects[0].source_content["source_verification"] == "SKIPPED"
        assert projects[1].source_content["confidence"] == 0.65
        # 兼容恢复该规则前已经保存的复合批次 hash；源文件 hash 仍保存在解析报告中。
        batch = session.scalar(select(ImportBatch))
        assert batch is not None
        batch.sha256 = "f" * 64

    duplicate = service.import_structured_topics(
        actor_name="老板测试",
        original_filename="换个名字也还是同一内容.docx",
        idempotency_key="structured-import-0002",
        topics=topics,
        warnings=[],
        schema_version="1.0",
        file_key=file_key,
        file_url=None,
    )
    assert duplicate["data"]["deduplicated"] is True
    assert duplicate["data"]["created_count"] == 0

    changed_extraction = service.import_structured_topics(
        actor_name="老板测试",
        original_filename="还是同一个原文件.docx",
        idempotency_key="structured-import-0003",
        topics=[
            StructuredTopicInput(
                title="同一文件重新提取的新任务",
                source_text="MCP 重新整理后的任务内容。",
                script=None,
                confidence=0.5,
                evidence=[],
            )
        ],
        warnings=[],
        schema_version="1.0",
        file_key=file_key,
        file_url=None,
    )
    assert changed_extraction["data"]["deduplicated"] is True
    assert changed_extraction["data"]["created_count"] == 0
    assert len(changed_extraction["data"]["tasks"]) == 2


def test_structured_import_trusts_mcp_content_without_source_verification(tmp_path) -> None:
    service = make_service(tmp_path)
    file_key = upload_file(
        service,
        actor_name="老板测试",
        role="BOSS",
        content=_arbitrary_document(),
    )

    imported = service.import_structured_topics(
        actor_name="老板测试",
        original_filename="任意格式.docx",
        idempotency_key="structured-trusted-0001",
        topics=[
            StructuredTopicInput(
                title="MCP 生成的任务",
                source_text="这段内容不要求能在原始 Word 文档中定位。",
                script="MCP 也可以直接提供生成后的口播稿。",
                confidence=0.2,
                evidence=["模型整理的定位说明"],
            )
        ],
        warnings=[],
        schema_version="1.0",
        file_key=file_key,
        file_url=None,
    )

    assert imported["data"]["created_count"] == 1
    assert imported["data"]["warnings"] == []
    with service.sessions() as session:
        project = session.scalar(select(ContentProject))
        assert project is not None
        assert project.source_content["source_text"].startswith("这段内容不要求")
        assert project.source_content["source_verification"] == "SKIPPED"
