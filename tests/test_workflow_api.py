"""内部工作流 API 端到端测试（规格 §3.1 权限双层校验的 API 层）。

通过 FastAPI TestClient 走真实 HTTP 路由：Token → 身份 → 角色 → 业务服务。
数据库使用临时 SQLite 文件，文件存储使用临时目录。
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient

BOSS = {"Authorization": "Bearer boss-token-at-least-16"}
EMP_A = {"Authorization": "Bearer employee-a-token-12345"}
EMP_B = {"Authorization": "Bearer employee-b-token-12345"}
UPLOAD = {"Authorization": "Bearer fixed-upload-token-at-least-16"}


@pytest.fixture
def api(tmp_path, monkeypatch):
    monkeypatch.setenv("APP_ENV", "test")
    monkeypatch.setenv("DATABASE_URL", f"sqlite+pysqlite:///{tmp_path / 'api.sqlite'}")
    monkeypatch.setenv("FILE_DATA_DIR", str(tmp_path / "files"))
    monkeypatch.setenv("PROBE_DATA_DIR", str(tmp_path / "probes"))
    monkeypatch.setenv("PUBLIC_BASE_URL", "http://testserver")
    monkeypatch.setenv("COMPANY_ID", "company-api")
    monkeypatch.setenv("MCP_BOSS_NAME", "老板测试")
    monkeypatch.setenv("MCP_BOSS_TOKEN", "boss-token-at-least-16")
    monkeypatch.setenv("FILE_UPLOAD_TOKEN", "fixed-upload-token-at-least-16")
    monkeypatch.setenv(
        "MCP_EMPLOYEES_JSON",
        '[{"name":"员工甲","token":"employee-a-token-12345","wecom_userid":"","active":true},'
        '{"name":"员工乙","token":"employee-b-token-12345","wecom_userid":"","active":true}]',
    )
    from workflow.config import get_settings

    get_settings.cache_clear()
    import workflow.apps.workflow_api as module
    from workflow.db.models import Base

    Base.metadata.create_all(module.engine)
    with TestClient(module.app) as client:
        yield client
    get_settings.cache_clear()


def upload_file(api, content: bytes) -> str:
    response = api.post(
        "/api/files/upload",
        headers={**UPLOAD, "Content-Type": "application/octet-stream"},
        content=content,
    )
    assert response.status_code == 200
    return response.json()["data"]["file_key"]


def test_auth_matrix_and_unified_error_structure(api) -> None:
    ok = api.get("/internal/t1/identity", headers=BOSS)
    assert ok.status_code == 200
    assert ok.json()["data"]["role"] == "BOSS"

    missing = api.get("/internal/t1/identity")
    assert missing.status_code == 401
    assert missing.json()["success"] is False
    assert missing.json()["error"]["code"] == "UNAUTHENTICATED"
    assert missing.json()["request_id"].startswith("req_")
    assert missing.json()["error"]["remediation"]

    invalid = api.get("/internal/t1/identity", headers={"Authorization": "Bearer nope-nope-nope"})
    assert invalid.status_code == 401
    assert invalid.json()["error"]["code"] == "UNAUTHENTICATED"

    # 员工调用老板端点 → 403 FORBIDDEN；老板调用员工端点同样拒绝。
    forbidden = api.post("/internal/tools/list-employees", headers=EMP_A, json={})
    assert forbidden.status_code == 403
    assert forbidden.json()["error"]["code"] == "FORBIDDEN"

    employee_reassign = api.post(
        "/internal/tools/change-task-assignee",
        headers=EMP_A,
        json={
            "task_no": "TASK-20260730-0001",
            "new_employee_id": "employee-b",
            "idempotency_key": "employee-reassign-forbidden-0001",
        },
    )
    assert employee_reassign.status_code == 403
    assert employee_reassign.json()["error"]["code"] == "FORBIDDEN"

    boss_on_employee = api.post("/internal/tools/list-my-tasks", headers=BOSS, json={})
    assert boss_on_employee.status_code == 403
    assert boss_on_employee.json()["error"]["code"] == "FORBIDDEN"

    bad_params = api.post("/internal/tools/import-topic-document", headers=BOSS, json={})
    assert bad_params.status_code == 400
    assert bad_params.json()["error"]["code"] == "INVALID_ARGUMENT"

    not_found = api.get("/files/does-not-exist-at-all")
    assert not_found.status_code == 404
    assert not_found.json()["error"]["code"] == "RESOURCE_NOT_FOUND"


def test_standalone_file_upload_api_contract(api) -> None:
    content = b"standalone-upload"
    uploaded = api.post(
        "/api/files/upload",
        headers={**UPLOAD, "Content-Type": "application/octet-stream"},
        content=content,
    )
    assert uploaded.status_code == 200
    assert uploaded.json()["data"]["file_key"]
    assert uploaded.json()["data"]["size_bytes"] == len(content)

    missing_auth = api.post(
        "/api/files/upload",
        headers={"Content-Type": "application/octet-stream"},
        content=content,
    )
    assert missing_auth.status_code == 401

    wrong_media_type = api.post(
        "/api/files/upload",
        headers={**UPLOAD, "Content-Type": "application/json"},
        content=content,
    )
    assert wrong_media_type.status_code == 400
    assert wrong_media_type.json()["error"]["code"] == "INVALID_ARGUMENT"

    empty = api.post(
        "/api/files/upload",
        headers={**UPLOAD, "Content-Type": "application/octet-stream"},
        content=b"",
    )
    assert empty.status_code == 400

    removed_mcp_upload_route = api.post(
        "/internal/tools/upload-file",
        headers=BOSS,
        json={"file_base64": "ZmlsZQ=="},
    )
    assert removed_mcp_upload_route.status_code == 404

    personal_token = api.post(
        "/api/files/upload",
        headers={**BOSS, "Content-Type": "application/octet-stream"},
        content=content,
    )
    assert personal_token.status_code == 401

    page = api.get("/file-upload")
    assert page.status_code == 200
    assert "不需要安装 Node.js 或 Python" in page.text
    assert page.text == Path("src/workflow/static/upload-file.html").read_text()

    preflight = api.options(
        "/api/files/upload",
        headers={
            "Origin": "null",
            "Access-Control-Request-Method": "POST",
            "Access-Control-Request-Headers": "authorization,content-type",
        },
    )
    assert preflight.status_code == 200
    assert preflight.headers["access-control-allow-origin"] == "null"


def test_structured_topic_import_api_accepts_non_template_word(api) -> None:
    document = Document()
    document.add_paragraph("这是一份没有固定标题样式的 ChatGPT 文档。")
    document.add_paragraph("企业开始用智能代理处理重复性内容工作。")
    buffer = BytesIO()
    document.save(buffer)
    file_key = upload_file(api, buffer.getvalue())
    body = {
        "original_filename": "自由格式.docx",
        "idempotency_key": "api-structured-0001",
        "file_key": file_key,
        "topics": [
            {
                "title": "智能代理进入内容工作流",
                "source_text": "企业开始用智能代理处理重复性内容工作。",
                "script": None,
                "confidence": 0.92,
                "evidence": ["智能代理处理重复性内容工作"],
            }
        ],
        "warnings": [],
        "schema_version": "1.0",
    }

    imported = api.post("/internal/tools/import-structured-topics", headers=BOSS, json=body)
    assert imported.status_code == 200
    assert imported.json()["data"]["import_mode"] == "WORKBUDDY_STRUCTURED"
    assert imported.json()["data"]["created_count"] == 1

    public_body = {
        **body,
        "idempotency_key": "api-structured-public-0001",
    }
    public_import = api.post(
        "/api/topics/import-structured",
        headers=BOSS,
        json=public_body,
    )
    assert public_import.status_code == 200
    assert public_import.json()["data"]["deduplicated"] is True
    assert public_import.json()["data"]["created_count"] == 0

    public_forbidden = api.post(
        "/api/topics/import-structured",
        headers=EMP_A,
        json=public_body,
    )
    assert public_forbidden.status_code == 403

    legacy_body = {**body, "content_base64": "ZmlsZQ=="}
    legacy_body.pop("file_key")
    legacy = api.post(
        "/internal/tools/import-structured-topics",
        headers=BOSS,
        json=legacy_body,
    )
    assert legacy.status_code == 400
    assert legacy.json()["error"]["code"] == "INVALID_ARGUMENT"

    forbidden = api.post("/internal/tools/import-structured-topics", headers=EMP_A, json=body)
    assert forbidden.status_code == 403


def test_full_topic_script_flow_through_internal_api(api) -> None:
    document = Path("docs/AI行业选题文档上传样例.docx").read_bytes()
    topic_file_key = upload_file(api, document)
    imported = api.post(
        "/internal/tools/import-topic-document",
        headers=BOSS,
        json={
            "original_filename": "选题.docx",
            "idempotency_key": "api-import-0001",
            "file_key": topic_file_key,
        },
    )
    assert imported.status_code == 200
    data = imported.json()["data"]
    assert data["created_count"] == 10
    assert data["parse_status"] == "COMPLETED"

    task = data["tasks"][0]
    owner_headers = EMP_A if task["assigned_employee_name"] == "员工甲" else EMP_B
    other_headers = EMP_B if owner_headers is EMP_A else EMP_A

    # 他人任务对员工不可见。
    invisible = api.post(
        "/internal/tools/get-my-task", headers=other_headers, json={"task_no": task["task_no"]}
    )
    assert invisible.status_code == 404
    assert invisible.json()["error"]["code"] == "RESOURCE_NOT_FOUND"

    mine = api.post("/internal/tools/list-my-tasks", headers=owner_headers, json={})
    assert any(item["task_no"] == task["task_no"] for item in mine.json()["data"])

    script_file_key = upload_file(api, "第一版".encode())
    submitted = api.post(
        "/internal/tools/submit-script-file",
        headers=owner_headers,
        json={
            "task_no": task["task_no"],
            "original_filename": "演播稿.txt",
            "idempotency_key": "api-submit-0001",
            "file_key": script_file_key,
            "note": "初稿",
        },
    )
    assert submitted.status_code == 200
    assert submitted.json()["data"]["version_no"] == 1

    pending = api.post("/internal/tools/list-pending-reviews", headers=BOSS, json={})
    assert any(item["task_no"] == task["task_no"] for item in pending.json()["data"])

    reviewed = api.post(
        "/internal/tools/review-script-submission",
        headers=BOSS,
        json={
            "task_no": task["task_no"],
            "decision": "APPROVED",
            "idempotency_key": "api-review-0001",
        },
    )
    assert reviewed.status_code == 200
    assert reviewed.json()["data"]["project_status"] == "WAITING_FOR_FILMING"

    # 相同幂等键重放返回首次结果，不产生重复审核。
    replay = api.post(
        "/internal/tools/review-script-submission",
        headers=BOSS,
        json={
            "task_no": task["task_no"],
            "decision": "APPROVED",
            "idempotency_key": "api-review-0001",
        },
    )
    assert replay.status_code == 200
    assert replay.json() == reviewed.json()
