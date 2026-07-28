from __future__ import annotations

from datetime import UTC, datetime
from io import BytesIO
from pathlib import Path
from zoneinfo import ZoneInfo

from docx import Document
from sqlalchemy import create_engine
from sqlalchemy.orm import Session
from sqlalchemy.pool import StaticPool

from workflow.db.models import ActorProfile, Base, Role, TaskAssignment
from workflow.t2.allocation import choose_employee, employee_loads
from workflow.t2.calendar import effective_started_at, week_start_for
from workflow.t2.parser import parse_topic_document


def test_real_topic_sample_parses_exactly_ten_topics() -> None:
    content = Path("docs/AI行业选题文档上传样例.docx").read_bytes()

    result = parse_topic_document(content)

    assert len(result.topics) == 10
    assert result.failures == []
    assert [topic.source_sequence for topic in result.topics] == list(range(1, 11))
    assert result.topics[0].title.startswith("芯片公司先给客户50亿美元")
    assert all(topic.script for topic in result.topics)
    assert all(
        "506字" not in topic.script and "505字" not in topic.script for topic in result.topics
    )


def _build_docx(blocks: list[tuple[str, str] | tuple[str, str, list[list[str]]]]) -> bytes:
    document = Document()
    for block in blocks:
        style, text = block[0], block[1]
        if style == "TABLE":
            rows = block[2]  # type: ignore[index]
            table = document.add_table(rows=len(rows), cols=len(rows[0]))
            for row, cells in zip(table.rows, rows, strict=True):
                for cell, value in zip(row.cells, cells, strict=True):
                    cell.text = value
        else:
            document.add_paragraph(text, style=style)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_missing_script_records_failure_detail() -> None:
    content = _build_docx(
        [
            ("Heading 1", "01｜有正文"),
            ("Normal", "主爆款标题：标题一"),
            ("Heading 2", "提词器口播正文（正文汉字数：10字）"),
            ("Normal", "这是口播正文。"),
            ("Heading 1", "02｜缺正文"),
            ("Normal", "主爆款标题：标题二"),
        ]
    )

    result = parse_topic_document(content)

    assert [topic.source_sequence for topic in result.topics] == [1]
    assert len(result.failures) == 1
    assert result.failures[0].source_sequence == 2
    assert result.failures[0].heading_title == "缺正文"
    assert "提词器口播正文" in result.failures[0].reason


def test_fields_inside_tables_are_parsed_in_document_order() -> None:
    content = _build_docx(
        [
            ("Heading 1", "01｜表格选题"),
            ("TABLE", "", [["主爆款标题", "表格里的标题"], ["所属领域", "算力"]]),
            ("Heading 2", "提词器口播正文（正文汉字数：8字）"),
            ("TABLE", "", [["第一段口播正文。"], ["第二段口播正文。"]]),
        ]
    )

    result = parse_topic_document(content)

    assert result.failures == []
    assert len(result.topics) == 1
    topic = result.topics[0]
    assert topic.title == "表格里的标题"
    assert "第一段口播正文。" in topic.script
    assert "第二段口播正文。" in topic.script


def test_empty_document_reports_generic_failure() -> None:
    content = _build_docx([("Normal", "没有任何选题的文档。")])

    result = parse_topic_document(content)

    assert result.topics == []
    assert len(result.failures) == 1
    assert "选题区块" in result.failures[0].reason


def test_unnumbered_heading_stops_topic_block() -> None:
    content = _build_docx(
        [
            ("Heading 1", "01｜唯一选题"),
            ("Normal", "主爆款标题：标题"),
            ("Heading 2", "提词器口播正文（正文汉字数：5字）"),
            ("Normal", "口播正文。"),
            ("Heading 1", "文末核验表"),
            ("Normal", "主爆款标题：不应生成任务"),
            ("Heading 2", "提词器口播正文（正文汉字数：5字）"),
            ("Normal", "不应被解析。"),
        ]
    )

    result = parse_topic_document(content)

    assert len(result.topics) == 1
    assert result.topics[0].source_sequence == 1


def test_effective_started_at_business_boundaries() -> None:
    zone = ZoneInfo("Asia/Shanghai")

    monday_early = datetime(2026, 7, 27, 8, 30, tzinfo=zone)
    monday_working = datetime(2026, 7, 27, 10, 0, tzinfo=zone)
    friday_late = datetime(2026, 7, 31, 18, 0, tzinfo=zone)
    saturday = datetime(2026, 8, 1, 12, 0, tzinfo=zone)

    assert effective_started_at(monday_early).hour == 9
    assert effective_started_at(monday_working) == monday_working
    assert effective_started_at(friday_late) == datetime(2026, 8, 3, 9, 0, tzinfo=zone)
    assert effective_started_at(saturday) == datetime(2026, 8, 3, 9, 0, tzinfo=zone)


def test_weekly_load_uses_append_only_deltas_and_minimum_choice() -> None:
    engine = create_engine(
        "sqlite+pysqlite:///:memory:",
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    Base.metadata.create_all(engine)
    now = datetime(2026, 7, 28, 2, 0, tzinfo=UTC)
    week_start = week_start_for(now)
    with Session(engine) as session:
        a = ActorProfile(
            company_id="company",
            display_name="员工甲",
            role=Role.EMPLOYEE,
            position="新媒体运营",
            active=True,
            token_sha256="a" * 64,
        )
        b = ActorProfile(
            company_id="company",
            display_name="员工乙",
            role=Role.EMPLOYEE,
            position="新媒体运营",
            active=True,
            token_sha256="b" * 64,
        )
        session.add_all([a, b])
        session.flush()
        session.add_all(
            [
                TaskAssignment(
                    company_id="company",
                    task_id="task-a",
                    assignee_id=a.id,
                    event_type="AUTO_ASSIGNED",
                    workload_delta=1,
                    work_week_start=week_start,
                    assigned_at=now,
                ),
                TaskAssignment(
                    company_id="company",
                    task_id="task-a",
                    assignee_id=a.id,
                    event_type="CANCEL_REVERSED",
                    workload_delta=-1,
                    work_week_start=week_start,
                    assigned_at=now,
                ),
                TaskAssignment(
                    company_id="company",
                    task_id="task-b",
                    assignee_id=b.id,
                    event_type="AUTO_ASSIGNED",
                    workload_delta=1,
                    work_week_start=week_start,
                    assigned_at=now,
                ),
            ]
        )
        session.commit()

        loads = employee_loads(session, "company", now, "Asia/Shanghai")
        assert {employee.display_name: count for employee, count in loads} == {
            "员工甲": 0,
            "员工乙": 1,
        }
        assert choose_employee(loads).display_name == "员工甲"
